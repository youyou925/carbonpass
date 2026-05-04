from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('比特币（BTC）行情分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('2025年4月13日')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

# 一、核心数据概览
doc.add_heading('一、核心数据概览', level=1)
items = [
    '当前价格：$XX,XXX',
    '24小时涨跌幅：+/- X.XX%',
    '24小时最低/最高：$XX,XXX – $XX,XXX',
    '24小时成交量：$XX Billion',
    '市场情绪指数：中性偏多（XX）'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# 二、技术面分析
doc.add_heading('二、技术面分析', level=1)
doc.add_paragraph('支撑位：$XX,XXX（前低支撑）')
doc.add_paragraph('阻力位：$XX,XXX（前期高点）')
doc.add_paragraph('均线系统：MA5 上穿 MA10，短期偏多')
doc.add_paragraph('MACD：即将形成金叉')
doc.add_paragraph('RSI：62（中性偏强区域）')

# 三、影响因素
doc.add_heading('三、今日重要影响因素', level=1)
doc.add_paragraph('1. 宏观事件：今晚 20:30 美国 CPI 数据公布', style='List Number')
doc.add_paragraph('2. 政策消息：美国比特币现货ETF 昨日净流入 $X 亿', style='List Number')
doc.add_paragraph('3. 链上数据：交易所 BTC 余额减少，巨鲸积累迹象明显', style='List Number')

# 四、操作建议
doc.add_heading('四、短期操作建议（仅供参考）', level=1)
doc.add_paragraph('• 多单策略：价格站稳 $XX,XXX，轻仓试多，目标 $XX,XXX')
doc.add_paragraph('• 空单策略：若跌破 $XX,XXX 且无法收回，可考虑反弹做空')
doc.add_paragraph('• 观望策略：CPI数据公布前，不建议重仓')

# 五、风险提示
doc.add_heading('五、风险提示', level=1)
doc.add_paragraph('本报告仅供分析参考，不构成投资建议。数字货币市场波动极大，请做好仓位管理与止损。')

# 保存文件
doc.save('比特币行情分析报告.docx')
print('报告生成成功：比特币行情分析报告.docx')
